import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def assign_group(val, template, colname='Description'):
    if pd.isna(val): return None
    val = str(val).strip().lower()
    for _, row in template.iterrows():
        desc = str(row[colname]).strip().lower()
        if val == desc:
            return row['Group']
    return None

def assign_range_group(value, template):
    try: value = float(value)
    except: return None
    for _, row in template.iterrows():
        if row['Min'] <= value <= row['Max']:
            return row['Group']
    return None

def make_rating_table(df):
    tbl = df[['Sr. No.', 'Product Name', 'Solubility_Group', 'Dose_Group', 'Toxicity_Group', 'Cleaning_Group']].copy()
    tbl.rename(columns={
        'Solubility_Group': 'Solubility',
        'Dose_Group': 'TDD (mg)',
        'Toxicity_Group': 'Toxicity',
        'Cleaning_Group': 'Hardest To Clean'
    }, inplace=True)
    tbl['Total'] = tbl['Solubility'] + tbl['TDD (mg)'] + tbl['Toxicity'] + tbl['Hardest To Clean']
    return tbl

def add_table_from_df(doc, df, title=None, style="Table Grid", max_cols=7):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(0,51,102)
        run.font.size = Pt(13)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    if len(df.columns) <= max_cols:
        _add_table(doc, df, style)
    else:
        col_chunks = [df.columns[i:i+max_cols] for i in range(0, len(df.columns), max_cols)]
        for idx, cols in enumerate(col_chunks):
            _add_table(doc, df[list(cols)], style, f"Part {idx+1}")
    doc.add_paragraph()
def _add_table(doc, df, style="Table Grid", part_text=None):
    if part_text:
        doc.add_paragraph(part_text).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table = doc.add_table(rows=1, cols=len(df.columns), style=style)
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def add_paragraph_justify(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return p

def generate_word_report(df, df_equip, solubility_template, dose_template, toxicity_template, cleaning_template,
                         rating_table, prev_worst_case, next_worst_case, 
                         maco_10ppm, maco_tdd, maco_ade, lowest_maco,
                         swab_surface, total_surface_area_with_margin, swab_limit,
                         df_rinse_limits):
    doc = Document()
    doc.add_heading("CLEANING VALIDATION PROTOCOL", 1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Date: .......................................").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p = doc.add_paragraph()
    run = p.add_run("1. Introduction")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0,51,102)
    add_paragraph_justify(doc,
        "This report documents the systematic evaluation of cleaning validation using MACO (Maximum Allowable Carryover), Swab, and Rinse limit calculations. The protocol is based on worst-case product selection, group rating assignment for solubility, dose, toxicity, and cleaning difficulty, and includes analytical and equipment considerations. All calculations are performed according to current industry guidelines and in compliance with regulatory expectations."
    )
    run = doc.add_paragraph()
    run = run.add_run("2. Objective")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0,51,102)
    add_paragraph_justify(doc,
        "To establish scientifically justified cleaning limits for shared equipment, based on product and process understanding, using group ratings, worst-case identification, and three MACO calculation methods (10 ppm, TDD, ADE/PDE)."
    )
    run = doc.add_paragraph()
    run = run.add_run("3. Product Group Ratings Table")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0,51,102)
    add_table_from_df(doc, rating_table, title="Product Group Ratings (by sum)")
    run = doc.add_paragraph()
    run = run.add_run("4. Product Details - Partwise Tables")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0,51,102)
    product_cols = [
        ['Sr. No.', 'Product Name', 'Solubility', 'Solubility_Group', 'Hardest To Clean', 'Cleaning_Group'],
        ['Sr. No.', 'Product Name', 'Min Batch Size (kg)', 'Max Batch Size (kg)', 'Min Dose (mg)', 'Dose_Group', 'ADE/PDE (µg/day)', 'Toxicity_Group', 'Worst_Case_Rating'],
        ['Sr. No.', 'Product Name', 'Swab Recovery %', 'LOD in ppm', 'LOQ in ppm', 'Swab Dilution in ml as per AMV', 'Swab Surface in M. Sq.']
    ]
    for i, cols in enumerate(product_cols):
        partdf = df[cols] if all(col in df.columns for col in cols) else df[[col for col in cols if col in df.columns]]
        add_table_from_df(doc, partdf, title=f"Product Details - Part {i+1}")
    run = doc.add_paragraph()
    run = run.add_run("5. Equipment Details")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0,51,102)
    add_table_from_df(doc, df_equip[['Eq. Name', 'Eq. ID', 'Product contact Surface Area (m2)', 'Used for', 'Cleaning   Procedure No.']], title="Equipment List")
    run = doc.add_paragraph()
    run = run.add_run("6. Group Definitions")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0,51,102)
    add_table_from_df(doc, solubility_template, title="Solubility Group Definition")
    add_table_from_df(doc, dose_template, title="Dose Group Definition")
    add_table_from_df(doc, toxicity_template, title="Toxicity Group Definition")
    add_table_from_df(doc, cleaning_template, title="Hardest To Clean Group Definition")
    run = doc.add_paragraph()
    run = run.add_run("7. Worst Case Product Selection")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0,51,102)
    add_paragraph_justify(doc, "Previous Worst Case (Highest Product of Group Multiplication Rating):")
    add_table_from_df(doc, prev_worst_case.to_frame().T[['Product Name', 'Worst_Case_Rating', 'Min Batch Size (kg)', 'ADE/PDE (µg/day)', 'Min Dose (mg)', 'Max Dose (mg)']])
    add_paragraph_justify(doc, "Next Worst Case (Lowest Min Batch Size / Max Dose ratio):")
    add_table_from_df(doc, next_worst_case.to_frame().T[['Product Name', 'BatchSize_Dose_Ratio', 'Min Batch Size (kg)', 'ADE/PDE (µg/day)', 'Min Dose (mg)', 'Max Dose (mg)']])
    run = doc.add_paragraph()
    run = run.add_run("8. MACO (Maximum Allowable Carryover) Calculations")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0,51,102)
    macodf = pd.DataFrame([
        ['10 ppm', f"{maco_10ppm:.4f} mg"],
        ['TDD', f"{maco_tdd:.4f} mg"],
        ['ADE/PDE', f"{maco_ade:.4f} mg"],
        ['Lowest MACO (used for limits)', f"{lowest_maco:.4f} mg"]
    ], columns=['MACO Calculation Method', 'Value (mg)'])
    add_table_from_df(doc, macodf, title="MACO Calculation Summary")
    run = doc.add_paragraph()
    run = run.add_run("9. Swab Limit Calculation")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0,51,102)
    swabdf = pd.DataFrame([{
        "Swab Surface in M. Sq.": swab_surface,
        "Total Equip Surface with 20% margin (m2)": round(total_surface_area_with_margin, 4),
        "Lowest MACO (mg)": round(lowest_maco, 4),
        "Swab Limit (mg)": round(swab_limit, 6)
    }])
    add_table_from_df(doc, swabdf, title="Swab Limit Summary")
    run = doc.add_paragraph()
    run = run.add_run("10. Rinse Limits and Volumes per Equipment")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0,51,102)
    add_table_from_df(doc, df_rinse_limits, title="Rinse Limits per Equipment")
    run = doc.add_paragraph()
    run = run.add_run("11. Conclusion & Summary")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0,51,102)
    add_paragraph_justify(doc, "The cleaning validation study demonstrates a robust, risk-based approach for setting carryover limits. The application of 20% margin to the total equipment surface area, group-based worst-case selection, and multiple MACO calculation approaches ensures regulatory compliance and patient safety. The lowest MACO value has been used for swab and rinse calculations. This protocol should be reviewed and approved prior to execution.")
    doc.add_paragraph(); doc.add_paragraph("Prepared by:  ___________________________")
    doc.add_paragraph("Reviewed by:  __________________________")
    doc.add_paragraph("Approved by:  __________________________")
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# --- Streamlit App ---
st.title("Cleaning Validation Protocol Generator")
st.write("Upload your Excel file and download calculation results (Excel & Word Report).")

uploaded_file = st.file_uploader("Upload Excel File", type=["xlsx"])
if uploaded_file:
    # Load data
    xls = pd.ExcelFile(uploaded_file)
    df = pd.read_excel(uploaded_file, sheet_name='Product Details')
    df_equip = pd.read_excel(uploaded_file, sheet_name='Equipment Details')
    solubility_template = pd.read_excel(uploaded_file, sheet_name='Solubility')
    dose_template = pd.read_excel(uploaded_file, sheet_name='Dose')
    toxicity_template = pd.read_excel(uploaded_file, sheet_name='Toxicity')
    cleaning_template = pd.read_excel(uploaded_file, sheet_name='Cleaning')
    # Assign groups
    df['Solubility_Group'] = df['Solubility'].apply(lambda x: assign_group(x, solubility_template))
    df['Dose_Group'] = df['Min Dose (mg)'].apply(lambda x: assign_range_group(x, dose_template))
    df['Toxicity_Group'] = df['ADE/PDE (µg/day)'].apply(lambda x: assign_range_group(x, toxicity_template))
    df['Cleaning_Group'] = df['Hardest To Clean'].apply(lambda x: assign_group(x, cleaning_template))
    rating_table = make_rating_table(df)
    # Worst case logic
    df['Worst_Case_Rating'] = (
        df['Solubility_Group'].astype(float) *
        df['Dose_Group'].astype(float) *
        df['Toxicity_Group'].astype(float) *
        df['Cleaning_Group'].astype(float)
    )
    df['BatchSize_Dose_Ratio'] = df['Min Batch Size (kg)'] / df['Max Dose (mg)']
    prev_worst_case = df.loc[df['Worst_Case_Rating'].idxmax()]
    next_worst_case = df.loc[df['BatchSize_Dose_Ratio'].idxmin()]
    min_batch_next_kg = next_worst_case['Min Batch Size (kg)']
    max_dose_next_mg = next_worst_case['Max Dose (mg)']
    min_dose_prev_mg = prev_worst_case['Min Dose (mg)']
    ade_prev_ug = prev_worst_case['ADE/PDE (µg/day)']
    ade_prev_mg = ade_prev_ug / 1000
    maco_10ppm = 0.00001 * min_batch_next_kg * 1e6 / max_dose_next_mg
    maco_tdd = min_dose_prev_mg * min_batch_next_kg * 1e6 / (max_dose_next_mg * 1000)
    maco_ade = ade_prev_mg * min_batch_next_kg * 1e6 / max_dose_next_mg
    lowest_maco = min(maco_10ppm, maco_tdd, maco_ade)
    total_surface_area = df_equip['Product contact Surface Area (m2)'].sum()
    total_surface_area_with_margin = total_surface_area * 1.2
    swab_surface = df['Swab Surface in M. Sq.'].iloc[0]
    swab_limit = lowest_maco * swab_surface / total_surface_area_with_margin
    rinse_limits = []
    for idx, row in df_equip.iterrows():
        eq_surface = row['Product contact Surface Area (m2)']
        rinse_limit = lowest_maco * eq_surface / total_surface_area_with_margin
        rinse_vol = rinse_limit / 10
        rinse_limits.append({
            'Eq. Name': row['Eq. Name'],
            'Eq. ID': row['Eq. ID'],
            'Surface Area (m2)': eq_surface,
            'Rinse Limit (mg)': round(rinse_limit, 6),
            'Rinse Volume (L)': round(rinse_vol, 6),
            'Rinse Volume (ml)': round(rinse_vol * 1000, 2)
        })
    df_rinse_limits = pd.DataFrame(rinse_limits)
    # Show tables in app
    st.subheader("Product Group Ratings Table")
    st.dataframe(rating_table)
    st.subheader("Worst Case Products")
    st.write("Previous Worst Case:", prev_worst_case['Product Name'])
    st.write("Next Worst Case:", next_worst_case['Product Name'])
    st.subheader("MACO Calculations (mg)")
    st.write(f"10 ppm: {maco_10ppm:.4f},  TDD: {maco_tdd:.4f},  ADE/PDE: {maco_ade:.4f},  Lowest Used: {lowest_maco:.4f}")
    st.subheader("Swab Limit (mg)")
    st.write(f"{swab_limit:.6f} mg per swab area ({swab_surface} m2)")
    st.subheader("Rinse Limits per Equipment")
    st.dataframe(df_rinse_limits)
    # Download Excel
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name="Products with Groups")
        df_equip.to_excel(writer, index=False, sheet_name="Equipment")
        df_rinse_limits.to_excel(writer, index=False, sheet_name="Rinse Limits")
        pd.DataFrame([{
            "MACO_10ppm (mg)": maco_10ppm,
            "MACO_TDD (mg)": maco_tdd,
            "MACO_ADE (mg)": maco_ade,
            "Lowest MACO Used (mg)": lowest_maco,
            "Swab Limit (mg)": swab_limit,
            "Total Equip Surface (m2)": total_surface_area,
            "Total Equip Surface with 20% margin (m2)": total_surface_area_with_margin,
            "Swab Surface Used (m2)": swab_surface
        }]).to_excel(writer, index=False, sheet_name="Summary")
    output.seek(0)
    st.download_button(
        label="Download Excel Results",
        data=output,
        file_name="cleaning_validation_with_margin.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    # Download Word
    docx_bytes = generate_word_report(
        df, df_equip, solubility_template, dose_template, toxicity_template, cleaning_template,
        rating_table, prev_worst_case, next_worst_case, 
        maco_10ppm, maco_tdd, maco_ade, lowest_maco,
        swab_surface, total_surface_area_with_margin, swab_limit,
        df_rinse_limits
    )
    st.download_button(
        label="Download Word Protocol Report",
        data=docx_bytes,
        file_name="Cleaning_Validation_Protocol_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
